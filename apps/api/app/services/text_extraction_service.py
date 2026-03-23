"""
Story 02.2: Text extraction from PDF and DOCX resume files.
Uses PyMuPDF for PDF and python-docx for DOCX.
"""
import io
import logging
from dataclasses import dataclass

logger = logging.getLogger(__name__)

MAX_TEXT_LENGTH = 50_000
SCANNED_TEXT_THRESHOLD = 100  # chars — below this we assume scanned PDF


@dataclass
class ExtractionResult:
    text: str
    is_scanned: bool
    page_count: int
    truncated: bool
    error: str | None = None


class TextExtractionService:
    """Extracts raw text from PDF and DOCX files."""

    def extract(self, file_bytes: bytes, content_type: str, file_name: str = "") -> ExtractionResult:
        """
        Dispatch to the correct extractor based on content type / file extension.
        Returns ExtractionResult with text and metadata.
        """
        ct = content_type.lower()
        name = file_name.lower()

        try:
            if ct == "application/pdf" or name.endswith(".pdf"):
                return self._extract_pdf(file_bytes)
            elif ct in (
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "application/msword",
            ) or name.endswith((".docx", ".doc")):
                return self._extract_docx(file_bytes)
            elif ct in ("application/zip", "application/x-zip-compressed") or name.endswith(".zip"):
                # Try to extract first PDF/DOCX from zip
                return self._extract_zip(file_bytes)
            else:
                return ExtractionResult(
                    text="", is_scanned=False, page_count=0, truncated=False,
                    error=f"Unsupported file type: {content_type}",
                )
        except Exception as exc:
            logger.exception("Text extraction failed for %s", file_name)
            return ExtractionResult(
                text="", is_scanned=False, page_count=0, truncated=False,
                error=f"Unable to extract text — file may be corrupted or password-protected: {exc}",
            )

    # ── PDF ───────────────────────────────────────────────────────────────────

    def _extract_pdf(self, file_bytes: bytes) -> ExtractionResult:
        import fitz  # PyMuPDF

        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        except Exception as exc:
            raise RuntimeError(f"Cannot open PDF: {exc}") from exc

        # Check for password protection
        if doc.needs_pass:
            raise RuntimeError("PDF is password-protected")

        pages = []
        for page in doc:
            pages.append(page.get_text())

        raw_text = "\n".join(pages).strip()
        is_scanned = len(raw_text) < SCANNED_TEXT_THRESHOLD
        truncated = len(raw_text) > MAX_TEXT_LENGTH

        if truncated:
            logger.warning("PDF text truncated at %d chars", MAX_TEXT_LENGTH)
            raw_text = raw_text[:MAX_TEXT_LENGTH]

        return ExtractionResult(
            text=raw_text,
            is_scanned=is_scanned,
            page_count=len(doc),
            truncated=truncated,
        )

    # ── DOCX ──────────────────────────────────────────────────────────────────

    def _extract_docx(self, file_bytes: bytes) -> ExtractionResult:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        try:
            doc = Document(io.BytesIO(file_bytes))
        except PackageNotFoundError as exc:
            raise RuntimeError(f"Cannot open DOCX: {exc}") from exc

        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables
        for table in doc.tables:
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_cells:
                    parts.append(" | ".join(row_cells))

        raw_text = "\n".join(parts).strip()
        truncated = len(raw_text) > MAX_TEXT_LENGTH

        if truncated:
            logger.warning("DOCX text truncated at %d chars", MAX_TEXT_LENGTH)
            raw_text = raw_text[:MAX_TEXT_LENGTH]

        return ExtractionResult(
            text=raw_text,
            is_scanned=False,
            page_count=0,  # DOCX has no page concept at extraction time
            truncated=truncated,
        )

    # ── ZIP ───────────────────────────────────────────────────────────────────

    def _extract_zip(self, file_bytes: bytes) -> ExtractionResult:
        import zipfile

        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = zf.namelist()
            # Try first PDF, then first DOCX
            for ext, ct in [(".pdf", "application/pdf"), (".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")]:
                for name in names:
                    if name.lower().endswith(ext) and not name.startswith("__MACOSX"):
                        inner_bytes = zf.read(name)
                        return self.extract(inner_bytes, ct, name)

        raise RuntimeError("No supported resume file found inside ZIP")
